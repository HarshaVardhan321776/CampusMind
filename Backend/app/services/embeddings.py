import os
import numpy as np
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import hashlib
import re
from langchain_core.embeddings import Embeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from app.core.config import settings

CHROMA_DIR = settings.CHROMA_DIR
COLLECTION_NAME = "campusmind_docs"


class LightweightHashEmbeddings(Embeddings):
    """Memory-safe, zero-PyTorch, deterministic 384-dim feature hashing embedding for lightweight deployments."""
    def __init__(self, dimension: int = 384):
        self.dimension = dimension

    def _embed_text(self, text: str) -> list[float]:
        if not text:
            return [0.0] * self.dimension
        vec = np.zeros(self.dimension, dtype=np.float32)
        words = re.findall(r"\b\w+\b", text.lower())
        for i, word in enumerate(words):
            # Token hashing
            h_val = int(hashlib.md5(word.encode("utf-8")).hexdigest(), 16)
            vec[h_val % self.dimension] += 1.0
            # Bigram hashing for local phrasing
            if i < len(words) - 1:
                bg = f"{word}_{words[i+1]}"
                h_bg = int(hashlib.md5(bg.encode("utf-8")).hexdigest(), 16)
                vec[h_bg % self.dimension] += 1.5
        norm = np.linalg.norm(vec)
        if norm > 0:
            vec = vec / norm
        return vec.tolist()

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._embed_text(t) for t in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._embed_text(text)


embedding_function = LightweightHashEmbeddings()

# Initialize OCR Engine for scanned/handwritten documents
_ocr_engine = None

def get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is None:
        try:
            from rapidocr_onnxruntime import RapidOCR
            _ocr_engine = RapidOCR()
        except Exception as e:
            print(f"[OCR Init Warning]: Could not initialize RapidOCR: {e}")
            _ocr_engine = False
    return _ocr_engine


def extract_pdf_with_ocr(file_path: str) -> list[Document]:
    """Extract text from scanned or image-based PDF using RapidOCR."""
    try:
        import pypdfium2 as pdfium
    except ImportError:
        print("[PDF Warning]: pypdfium2 not installed.")
        return []

    engine = get_ocr_engine()
    if not engine:
        return []

    documents = []
    try:
        pdf = pdfium.PdfDocument(file_path)
        for idx in range(len(pdf)):
            page = pdf[idx]
            # Render page to bitmap image
            bitmap = page.render(scale=1.5)
            pil_image = bitmap.to_pil()
            img_np = np.array(pil_image)

            ocr_res, _ = engine(img_np)
            if ocr_res:
                page_lines = [line[1] for line in ocr_res if line[1].strip()]
                page_text = "\n".join(page_lines)
                if len(page_text.strip()) > 10:
                    documents.append(
                        Document(
                            page_content=page_text.strip(),
                            metadata={"source": file_path, "page": idx + 1}
                        )
                    )
    except Exception as e:
        print(f"[OCR Extraction Error] on {file_path}: {e}")

    return documents


def load_document(file_path: str, file_type: str) -> list[Document]:
    """Load a PDF or DOCX file with automatic fallback to OCR for scanned documents."""
    if file_type == "docx":
        # 1. Try Docx2txtLoader
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(file_path)
            docs = loader.load()
            if docs and sum(len(d.page_content.strip()) for d in docs) > 0:
                return docs
        except Exception as e:
            print(f"[Docx2txtLoader Warning] on {file_path}: {e}")

        # 2. Native python-docx fallback (handles paragraphs and tables)
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            text = "\n\n".join(full_text)
            if text.strip():
                return [Document(page_content=text.strip(), metadata={"source": file_path})]
        except Exception as e:
            print(f"[python-docx Error] on {file_path}: {e}")
            raise RuntimeError(f"Could not extract text from docx file: {e}")

        return []

    elif file_type == "pdf":
        # 1. Attempt digital text extraction with PyPDFLoader
        digital_docs = []
        try:
            loader = PyPDFLoader(file_path)
            digital_docs = loader.load()
        except Exception as e:
            print(f"[PyPDFLoader Warning] on {file_path}: {e}")

        # Check quality of extracted digital text
        total_text_length = sum(len(doc.page_content.strip()) for doc in digital_docs)
        num_pages = len(digital_docs) if digital_docs else 1
        avg_chars_per_page = total_text_length / max(1, num_pages)

        # If digital text is missing, very sparse (< 40 chars/page), or empty, trigger OCR
        if total_text_length < 100 or avg_chars_per_page < 40:
            print(f"[CampusMind] Scanned/Image PDF detected for '{os.path.basename(file_path)}' (only {total_text_length} digital chars found). Running OCR pipeline...")
            ocr_docs = extract_pdf_with_ocr(file_path)
            if ocr_docs and sum(len(d.page_content) for d in ocr_docs) > total_text_length:
                print(f"[CampusMind] OCR successfully extracted {len(ocr_docs)} pages with {sum(len(d.page_content) for d in ocr_docs)} characters!")
                return ocr_docs

        # Ensure page numbers are 1-indexed for PyPDFLoader (which uses 0-indexed page numbers)
        for doc in digital_docs:
            if "page" in doc.metadata and isinstance(doc.metadata["page"], int):
                # If 0-indexed, convert to 1-indexed
                doc.metadata["page"] = doc.metadata["page"] + 1

        return digital_docs

    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def chunk_documents(documents: list[Document], chunk_size: int = 1200, chunk_overlap: int = 250):
    """Split loaded documents into contextual chunks preserving code blocks, tables, and definitions."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(documents)


def delete_document_embeddings(document_id: int, user_id: int | None = None):
    """Delete all vector embeddings associated with a given document_id from ChromaDB."""
    try:
        import chromadb
        client = chromadb.PersistentClient(path=CHROMA_DIR)
        collection = client.get_or_create_collection(COLLECTION_NAME)
        
        # Chroma where filter
        if user_id is not None:
            where_clause = {"$and": [{"document_id": int(document_id)}, {"user_id": int(user_id)}]}
        else:
            where_clause = {"document_id": int(document_id)}
            
        existing = collection.get(where=where_clause)
        if existing and existing.get("ids"):
            collection.delete(ids=existing["ids"])
            print(f"[CampusMind] Purged {len(existing['ids'])} vector chunks for doc #{document_id} from ChromaDB.")
        return True
    except Exception as e:
        print(f"[CampusMind Chroma Delete Error] for doc #{document_id}: {e}")
        return False


def embed_and_store(chunks: list[Document], source_filename: str, document_id: int, user_id: int):
    """Generate embeddings for chunks and store them in ChromaDB with complete user and document metadata."""
    if not chunks:
        print(f"[Embed Warning]: No text chunks to embed for {source_filename}")
        return 0

    # Purge any previous embeddings for this document ID to prevent duplicate chunk buildup
    delete_document_embeddings(document_id, user_id)

    for chunk in chunks:
        chunk.metadata["source"] = source_filename
        chunk.metadata["document_name"] = source_filename
        chunk.metadata["document_id"] = int(document_id)
        chunk.metadata["user_id"] = int(user_id)
        if "page" in chunk.metadata:
            try:
                chunk.metadata["page"] = int(chunk.metadata["page"])
            except Exception:
                pass

    vectorstore = Chroma(
        collection_name=COLLECTION_NAME,
        embedding_function=embedding_function,
        persist_directory=CHROMA_DIR,
    )
    vectorstore.add_documents(chunks)
    vectorstore.persist()

    return len(chunks)


def process_document(file_path: str, file_type: str, source_filename: str, document_id: int, user_id: int) -> int:
    """Full pipeline: load -> chunk -> embed -> store with user scoping. Returns number of chunks created."""
    documents = load_document(file_path, file_type)
    chunks = chunk_documents(documents)
    num_chunks = embed_and_store(chunks, source_filename, document_id, user_id)
    return num_chunks